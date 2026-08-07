"""
AI Resume Parser Module for ResumeAI Pro.
Handles PDF, DOCX, and text resume parsing, entity extraction, and section analysis.
"""

import os
import re
import io
from typing import Dict, List, Any, Union, BinaryIO, Optional

import PyPDF2
import docx

try:
    from resumeai_pro.utils.nlp_utils import (
        clean_text,
        extract_contact_info,
        extract_sections,
        extract_skills,
        extract_education,
        extract_years_experience
    )
except ImportError:
    from utils.nlp_utils import (
        clean_text,
        extract_contact_info,
        extract_sections,
        extract_skills,
        extract_education,
        extract_years_experience
    )


class ResumeParser:
    """
    Parses resume documents (PDF, DOCX, TXT) and extracts structured attributes.
    """

    def parse_pdf(self, file: Union[str, BinaryIO, io.BytesIO]) -> str:
        """
        Extract text from a PDF document using PyPDF2.
        
        :param file: File path string or binary file-like object (e.g. BytesIO)
        :return: Extracted and normalized text string
        """
        text_content = []
        file_stream = file

        # Open file if path string provided
        if isinstance(file, str):
            if not os.path.exists(file):
                raise FileNotFoundError(f"PDF file not found at path: {file}")
            file_stream = open(file, "rb")

        try:
            # Ensure stream position is at start if file-like object
            if hasattr(file_stream, "seek"):
                file_stream.seek(0)

            reader = PyPDF2.PdfReader(file_stream)
            num_pages = len(reader.pages)

            for page_idx in range(num_pages):
                page = reader.pages[page_idx]
                page_text = page.extract_text() or ""
                
                # Multi-column PDF handling: clean up broken line breaks within column blocks
                lines = page_text.splitlines()
                reconstructed_lines = []
                for line in lines:
                    line_str = line.strip()
                    if line_str:
                        reconstructed_lines.append(line_str)

                text_content.append("\n".join(reconstructed_lines))

            full_text = "\n\n".join(text_content)
            return clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF document: {str(e)}") from e
        finally:
            if isinstance(file, str) and hasattr(file_stream, "close"):
                file_stream.close()

    def parse_docx(self, file: Union[str, BinaryIO, io.BytesIO]) -> str:
        """
        Extract text from DOCX document using python-docx.
        Extracts both body paragraphs and table cell contents.
        
        :param file: File path string or binary file-like object
        :return: Extracted text string
        """
        try:
            if hasattr(file, "seek"):
                file.seek(0)

            doc = docx.Document(file)
            extracted_parts = []

            # Extract body paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_parts.append(paragraph.text.strip())

            # Extract table content (resumes often use tables for layouts)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        extracted_parts.append(" | ".join(row_data))

            full_text = "\n".join(extracted_parts)
            return clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX document: {str(e)}") from e

    def parse(self, file_path_or_file: Union[str, BinaryIO, io.BytesIO], file_type: Optional[str] = None) -> str:
        """
        Dispatcher method to extract text based on file type or extension.
        
        :param file_path_or_file: File path string or binary file object
        :param file_type: 'pdf', 'docx', or 'txt'. Inferred from path if None.
        :return: Extracted text
        """
        if file_type is None and isinstance(file_path_or_file, str):
            ext = os.path.splitext(file_path_or_file)[1].lower()
            if ext == ".pdf":
                file_type = "pdf"
            elif ext in [".docx", ".doc"]:
                file_type = "docx"
            elif ext == ".txt":
                file_type = "txt"
            else:
                raise ValueError(f"Unsupported file extension '{ext}'. Specify file_type explicitly.")

        if not file_type:
            file_type = "pdf"  # Default assumption if ambiguous stream

        file_type = file_type.lower().strip(".")

        if file_type == "pdf":
            return self.parse_pdf(file_path_or_file)
        elif file_type in ["docx", "doc"]:
            return self.parse_docx(file_path_or_file)
        elif file_type == "txt":
            if isinstance(file_path_or_file, str):
                with open(file_path_or_file, "r", encoding="utf-8", errors="ignore") as f:
                    return clean_text(f.read())
            elif hasattr(file_path_or_file, "read"):
                if hasattr(file_path_or_file, "seek"):
                    file_path_or_file.seek(0)
                raw_bytes = file_path_or_file.read()
                if isinstance(raw_bytes, str):
                    return clean_text(raw_bytes)
                return clean_text(raw_bytes.decode("utf-8", errors="ignore"))
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def extract_all(self, text: str) -> Dict[str, Any]:
        """
        Parse raw resume text into structured components: contact, sections, skills,
        education, experience, projects, certifications.
        
        :param text: Clean resume text string
        :return: Dictionary of extracted structured features
        """
        cleaned = clean_text(text)
        sections = extract_sections(cleaned)
        contact = extract_contact_info(cleaned)
        skills = extract_skills(cleaned)
        education = extract_education(sections.get("education", "") or cleaned)

        # Experience processing
        exp_text = sections.get("experience", "")
        years_exp = extract_years_experience(cleaned)
        experience = self._parse_experience_bullets(exp_text) if exp_text else []

        # Projects processing
        proj_text = sections.get("projects", "")
        projects = self._parse_projects(proj_text) if proj_text else []

        # Certifications processing
        cert_text = sections.get("certifications", "")
        certifications = self._parse_certifications(cert_text) if cert_text else []

        return {
            "contact": contact,
            "sections": sections,
            "skills": skills,
            "education": education,
            "experience": {
                "total_years": years_exp,
                "entries": experience,
                "raw_text": exp_text
            },
            "projects": projects,
            "certifications": certifications,
            "raw_text": cleaned
        }

    def _parse_experience_bullets(self, exp_text: str) -> List[Dict[str, Any]]:
        """Helper to break down experience section into structured items."""
        entries = []
        lines = [line.strip() for line in exp_text.splitlines() if line.strip()]
        
        current_entry = {"title_company": "", "bullets": []}
        for line in lines:
            if line.startswith(("-", "*", "•")):
                bullet = re.sub(r'^[-\*•]\s*', '', line)
                current_entry["bullets"].append(bullet)
            else:
                if current_entry["bullets"] or current_entry["title_company"]:
                    entries.append(current_entry)
                    current_entry = {"title_company": line, "bullets": []}
                else:
                    current_entry["title_company"] = line

        if current_entry["bullets"] or current_entry["title_company"]:
            entries.append(current_entry)

        return entries

    def _parse_projects(self, proj_text: str) -> List[Dict[str, Any]]:
        """Helper to structure project entries."""
        projects = []
        lines = [line.strip() for line in proj_text.splitlines() if line.strip()]
        for line in lines:
            clean_line = re.sub(r'^[-\*•]\s*', '', line)
            projects.append({"description": clean_line})
        return projects

    def _parse_certifications(self, cert_text: str) -> List[str]:
        """Helper to extract certification bullet items."""
        certs = []
        lines = [line.strip() for line in cert_text.splitlines() if line.strip()]
        for line in lines:
            clean_cert = re.sub(r'^[-\*•]\s*', '', line)
            certs.append(clean_cert)
        return certs
